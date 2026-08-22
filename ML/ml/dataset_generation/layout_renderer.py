"""
Layout Renderer: Renders document content dicts into PDF, DOCX, and PNG files.
Uses ReportLab for PDF, python-docx for DOCX, and Pillow/PDF conversion for PNG.
"""
import os
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import docx

def render_to_pdf(content: dict, output_path: str):
    """Renders structured document content to a PDF file using ReportLab."""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = SimpleDocTemplate(output_path, pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#1E3A8A'),
        spaceAfter=12
    )

    header_style = ParagraphStyle(
        'DocHeader',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#374151')
    )

    elements = []
    
    # Document Header Title
    doc_type_title = content.get("document_type", "BUSINESS DOCUMENT").replace("_", " ")
    elements.append(Paragraph(f"<b>{doc_type_title}</b>", title_style))
    elements.append(Spacer(1, 10))

    # Metadata block
    meta_text = f"<b>Doc ID:</b> {content.get('document_id', '')}<br/>"
    meta_text += f"<b>Date:</b> {content.get('date', '')}<br/>"
    meta_text += f"<b>Company:</b> {content.get('company_name', '')}<br/>"
    meta_text += f"<b>Counterparty:</b> {content.get('counterparty_name', '')}<br/>"
    meta_text += f"<b>Currency:</b> {content.get('currency', 'USD')}<br/>"
    
    if "payment_terms" in content:
        meta_text += f"<b>Payment Terms:</b> {content['payment_terms']}<br/>"
        
    elements.append(Paragraph(meta_text, header_style))
    elements.append(Spacer(1, 15))

    # Freeform body paragraphs if present
    if "body_text" in content and content["body_text"]:
        elements.append(Paragraph(content["body_text"], styles['Normal']))
        elements.append(Spacer(1, 15))

    # Table of line items if present
    items = content.get("items", [])
    if items:
        table_data = [["Description", "Qty", "Unit Price", "Total"]]
        for item in items:
            table_data.append([
                str(item.get("description", "")),
                str(item.get("quantity", "")),
                f"{item.get('unit_price', 0.0):.2f}",
                f"{item.get('total_price', 0.0):.2f}"
            ])
        
        # Subtotal / Total rows
        if "total_amount" in content:
            table_data.append(["", "", "<b>TOTAL:</b>", f"<b>{content['total_amount']:.2f}</b>"])

        t = Table(table_data, colWidths=[240, 60, 100, 100])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E5E7EB')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#111827')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#D1D5DB')),
        ]))
        elements.append(t)

    doc.build(elements)

def render_to_docx(content: dict, output_path: str):
    """Renders document content into a DOCX file using python-docx."""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = docx.Document()
    
    doc_type = content.get("document_type", "BUSINESS DOCUMENT").replace("_", " ")
    doc.add_heading(doc_type, level=0)

    p = doc.add_paragraph()
    p.add_run(f"Document ID: {content.get('document_id', '')}\n").bold = True
    p.add_run(f"Date: {content.get('date', '')}\n")
    p.add_run(f"Company: {content.get('company_name', '')}\n")
    p.add_run(f"Counterparty: {content.get('counterparty_name', '')}\n")
    p.add_run(f"Currency: {content.get('currency', '')}\n")

    if "body_text" in content and content["body_text"]:
        doc.add_paragraph(content["body_text"])

    items = content.get("items", [])
    if items:
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Description'
        hdr_cells[1].text = 'Qty'
        hdr_cells[2].text = 'Unit Price'
        hdr_cells[3].text = 'Total'

        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get("description", ""))
            row_cells[1].text = str(item.get("quantity", ""))
            row_cells[2].text = f"{item.get('unit_price', 0.0):.2f}"
            row_cells[3].text = f"{item.get('total_price', 0.0):.2f}"

        if "total_amount" in content:
            row_cells = table.add_row().cells
            row_cells[2].text = "TOTAL:"
            row_cells[3].text = f"{content['total_amount']:.2f}"

    doc.save(output_path)

def render_to_png(content: dict, output_path: str):
    """Renders document preview image using PIL."""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    width, height = 800, 1050
    img = Image.new('RGB', (width, height), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    # Simple clean text drawing
    doc_type = content.get("document_type", "BUSINESS DOCUMENT").replace("_", " ")
    draw.text((40, 40), doc_type, fill=(30, 58, 138))
    
    y = 80
    lines = [
        f"Document ID: {content.get('document_id', '')}",
        f"Date: {content.get('date', '')}",
        f"Company: {content.get('company_name', '')}",
        f"Counterparty: {content.get('counterparty_name', '')}",
        f"Currency: {content.get('currency', 'USD')}"
    ]
    for line in lines:
        draw.text((40, y), line, fill=(55, 65, 81))
        y += 25

    if "body_text" in content and content["body_text"]:
        y += 10
        draw.text((40, y), content["body_text"][:200] + "...", fill=(0, 0, 0))
        y += 40

    items = content.get("items", [])
    if items:
        y += 20
        draw.rectangle([40, y, 760, y + 25], fill=(229, 231, 235))
        draw.text((50, y + 5), "Description", fill=(0, 0, 0))
        draw.text((450, y + 5), "Qty", fill=(0, 0, 0))
        draw.text((550, y + 5), "Unit Price", fill=(0, 0, 0))
        draw.text((650, y + 5), "Total", fill=(0, 0, 0))
        y += 30

        for item in items:
            draw.text((50, y), str(item.get("description", ""))[:40], fill=(55, 65, 81))
            draw.text((450, y), str(item.get("quantity", "")), fill=(55, 65, 81))
            draw.text((550, y), f"{item.get('unit_price', 0.0):.2f}", fill=(55, 65, 81))
            draw.text((650, y), f"{item.get('total_price', 0.0):.2f}", fill=(55, 65, 81))
            y += 25

        if "total_amount" in content:
            y += 10
            draw.text((550, y), "TOTAL:", fill=(0, 0, 0))
            draw.text((650, y), f"{content['total_amount']:.2f}", fill=(0, 0, 0))

    img.save(output_path)
