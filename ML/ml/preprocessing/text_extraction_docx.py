"""
DOCX Text Extraction using python-docx.
"""
import docx

def extract_text_from_docx(docx_path: str) -> dict:
    """Extracts paragraph text and tables from a DOCX document."""
    doc = docx.Document(docx_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_str = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_str:
                table_texts.append(row_str)

    full_text = "\n".join(paragraphs + table_texts)
    return {
        "text": full_text,
        "page_count": 1,
        "paragraph_count": len(paragraphs),
        "table_count": len(doc.tables)
    }
