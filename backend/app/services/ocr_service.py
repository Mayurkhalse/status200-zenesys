import os
from typing import Dict, Any
import pymupdf as fitz  # PyMuPDF
import docx

class OCRService:
    def extract_text_from_bytes(self, file_bytes: bytes, mime_type: str, filename: str) -> Dict[str, Any]:
        """Routes file bytes based on MIME type / extension to appropriate parser."""
        ext = os.path.splitext(filename)[1].lower()

        if mime_type == "application/pdf" or ext == ".pdf":
            return self._parse_pdf(file_bytes)
        elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"] or ext == ".docx":
            return self._parse_docx(file_bytes)
        elif mime_type.startswith("image/") or ext in [".png", ".jpg", ".jpeg"]:
            return self._parse_image(file_bytes)
        else:
            # Plain text fallback
            try:
                text = file_bytes.decode('utf-8', errors='ignore')
                return {"text": text, "method": "text_fallback"}
            except Exception:
                return {"text": "", "method": "unknown"}

    def _parse_pdf(self, file_bytes: bytes) -> Dict[str, Any]:
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            extracted_text = []
            for page in doc:
                text = page.get_text()
                if text.strip():
                    extracted_text.append(text)
            full_text = "\n".join(extracted_text)
            return {"text": full_text, "method": "pymupdf", "page_count": len(doc)}
        except Exception as e:
            print(f"PyMuPDF parsing note: {e}")
            try:
                text = file_bytes.decode('utf-8', errors='ignore')
                return {"text": text, "method": "pdf_text_fallback"}
            except Exception:
                return {"text": "", "method": "pymupdf_failed"}

    def _parse_docx(self, file_bytes: bytes) -> Dict[str, Any]:
        try:
            import io
            doc = docx.Document(io.BytesIO(file_bytes))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
            return {"text": "\n".join(full_text), "method": "python-docx"}
        except Exception as e:
            print(f"python-docx parsing failed: {e}")
            return {"text": "", "method": "python-docx_failed"}

    def _parse_image(self, file_bytes: bytes) -> Dict[str, Any]:
        # PaddleOCR or image OCR fallback
        try:
            from paddleocr import PaddleOCR
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name
            
            ocr = PaddleOCR(use_angle_cls=True, lang='en', show_log=False)
            result = ocr.ocr(tmp_path, cls=True)
            os.remove(tmp_path)

            text_lines = []
            if result and result[0]:
                for line in result[0]:
                    text_lines.append(line[1][0])
            return {"text": "\n".join(text_lines), "method": "paddleocr"}
        except Exception as e:
            print(f"PaddleOCR fallback note: {e}")
            return {"text": "[Scanned Image Content]", "method": "image_ocr_fallback"}

ocr_service = OCRService()
